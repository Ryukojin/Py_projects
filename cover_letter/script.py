#https://code.visualstudio.com/docs/python/python-tutorial


import docx
import re

letter = docx.Document("Cover_letter.docx")

#Date
print(letter.paragraphs[6].text)
#Company name
name = letter.paragraphs[8].text
#Street Address
print(letter.paragraphs[9].text)
#Main address
print(letter.paragraphs[10].text)

#Para 1
print(letter.paragraphs[13].text)
#Last para
print(letter.paragraphs[17].text)

new_name = 'Microsoft'

phrase_1 = letter.paragraphs[13].text
re.sub(r"(^.*)CareSuper*(.*$)", r"\1Microsoft\2", phrase_1)

sentence = "I am applying for my candidature at CareSuper. Applying to CareSuper would expand my ability to solve real problems."
re.sub(r"(^.*)CareSuper*(.*$)", r"\1Microsoft\2", phrase_1)
re.sub(r"CareSuper", r"Microsoft", phrase_1)